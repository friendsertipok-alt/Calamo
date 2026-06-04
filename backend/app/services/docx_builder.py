from __future__ import annotations
"""
Calamo — DOCX Builder
Сборка документа Word по ГОСТу.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import re
import copy
import logging
import latex2mathml.converter
import mathml2omml
from datetime import datetime

logger = logging.getLogger("DocxBuilder")

from app.schemas.order import SourceItem, PaperOutline


class DocxBuilder:
    """Сборщик .docx документа по ГОСТу."""

    def __init__(self):
        self.doc = Document()
        self._footnote_id = 0  # счётчик сносок
        self.equation_counter = 0  # счётчик формул
        self.sources_dict: dict = {}
        self._setup_styles()
        self._setup_page_layout()
        self._setup_footnotes_part()
        self._setup_footnote_styles()
        # self._enable_update_fields()  # Отключено: нарушает XSD-последовательность w:settings

    def _enable_update_fields(self):
        """Включает флаг принудительного обновления полей (TOC) при открытии документа."""
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        element = parse_xml(r'<w:updateFields %s w:val="true"/>' % nsdecls('w'))
        self.doc.settings.element.append(element)

    def _setup_page_layout(self):
        """Настройка полей страницы по ГОСТу."""
        section = self.doc.sections[0]
        section.page_width = Mm(210)   # A4
        section.page_height = Mm(297)
        section.left_margin = Mm(30)   # ГОСТ: лево 30мм
        section.right_margin = Mm(10)  # ГОСТ: право 10мм
        section.top_margin = Mm(20)    # ГОСТ: верх 20мм
        section.bottom_margin = Mm(20) # ГОСТ: низ 20мм
        # Нижний колонтитул с номером страницы
        self._add_page_number_footer(section)

    def _add_page_number_footer(self, section):
        """Вставка автоматической нумерации страниц через XML."""
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)

        # Run с XML-полем для номера страницы
        run = p.add_run()
        fldChar_begin = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        run._element.append(fldChar_begin)

        run2 = p.add_run()
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
        run2._element.append(instrText)

        run3 = p.add_run()
        fldChar_end = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        run3._element.append(fldChar_end)

        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    def _setup_footnotes_part(self):
        """Создаём /word/footnotes.xml если его нет в документе."""
        try:
            pkg = self.doc.part
            # Пробуем получить уже существующий part сносок
            rels = pkg.rels
            self._footnotes_part = None
            for rel in rels.values():
                if 'footnotes' in rel.reltype.lower():
                    self._footnotes_part = rel.target_part
                    break

            if self._footnotes_part is None:
                # Создаём footnotes.xml с двумя обязательными элементами (separator, continuation)
                W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                nsmap = {'w': W, 'r': R}
                root = etree.Element('{%s}footnotes' % W, nsmap=nsmap)

                for ftype, fid in [('separator', -1), ('continuationSeparator', 0)]:
                    fn = etree.SubElement(root, '{%s}footnote' % W)
                    fn.set('{%s}type' % W, ftype)
                    fn.set('{%s}id' % W, str(fid))
                    fp = etree.SubElement(fn, '{%s}p' % W)
                    fr = etree.SubElement(fp, '{%s}r' % W)
                    if ftype == 'separator':
                        etree.SubElement(fr, '{%s}separator' % W)
                    else:
                        etree.SubElement(fr, '{%s}continuationSeparator' % W)

                import docx.opc.part as opc_part
                from docx.opc.packuri import PackURI
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                import docx.oxml as dxml

                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
                partname = PackURI('/word/footnotes.xml')
                blob = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
                self._footnotes_part = opc_part.Part(partname, content_type, blob, pkg.package)
                pkg.relate_to(self._footnotes_part,
                              'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')
        except Exception as e:
            # Если что-то пошло не так, отключаем настоящие сноски
            self._footnotes_part = None

    def _setup_footnote_styles(self):
        """Создаём стили FootnoteText и FootnoteReference максимально близко к системным Word."""
        styles = self.doc.styles

        # FootnoteText (Текст сноски)
        if 'CalamoFootnoteText' not in [s.name for s in styles]:
            try:
                ft = styles.add_style('CalamoFootnoteText', 1)
                ft.font.name = 'Times New Roman'
                ft.font.size = Pt(10)
                ft.paragraph_format.first_line_indent = Cm(1.25)
                ft.paragraph_format.left_indent = Cm(0)
                ft.paragraph_format.space_after = Pt(0)
                ft.paragraph_format.line_spacing = 1.0
            except Exception:
                pass

        # FootnoteReference (Знак сноски)
        if 'CalamoFootnoteReference' not in [s.name for s in styles]:
            try:
                fr = styles.add_style('CalamoFootnoteReference', 2)
                fr.font.name = 'Times New Roman'
                fr.font.size = Pt(10)
                fr.font.superscript = True
            except Exception:
                pass

    def _setup_styles(self):
        """Настройка стилей по ГОСТу."""
        # Основной текст
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(14)
        font.color.rgb = RGBColor(0, 0, 0)
        
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5 интервал
        pf.first_line_indent = Cm(1.25)  # Абзацный отступ
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # По ширине
        pf.keep_with_next = False
        pf.keep_together = False
        pf.widow_control = True  # Запрет висячих строк

        # Заголовок 1 (главы)
        h1 = self.doc.styles["Heading 1"]
        h1.font.name = "Times New Roman"
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.paragraph_format.space_before = Pt(0)
        h1.paragraph_format.space_after = Pt(0)
        h1.paragraph_format.first_line_indent = Cm(0)
        h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # Заголовок 2 (параграфы)
        h2 = self.doc.styles["Heading 2"]
        h2.font.name = "Times New Roman"
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.space_before = Pt(0)
        h2.paragraph_format.space_after = Pt(0)
        h2.paragraph_format.first_line_indent = Cm(1.25)
        h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # Стиль для ссылок
        if 'CalamoHyperlink' not in self.doc.styles:
            hyperlink = self.doc.styles.add_style('CalamoHyperlink', 2) # 2 = character style
            hyperlink.font.name = "Times New Roman"
            hyperlink.font.color.rgb = RGBColor(0, 0, 255) # Синий
            hyperlink.font.underline = True

        # Стили для оглавления (Calamo TOC)
        for i in range(1, 4):
            style_name = f'Calamo TOC {i}'
            try:
                if style_name in self.doc.styles:
                    s = self.doc.styles[style_name]
                else:
                    s = self.doc.styles.add_style(style_name, 1)
                
                # Изолируем стиль от Normal, чтобы настройки не "слетали" при наследовании
                s.base_style = None
                
                s.font.name = "Times New Roman"
                s.font.size = Pt(14)
                s.font.bold = False
                s.font.color.rgb = RGBColor(0, 0, 0)
                
                # Через python-docx API (дублируем)
                s.paragraph_format.left_indent = Cm(0)
                s.paragraph_format.right_indent = Cm(0)
                s.paragraph_format.first_line_indent = Cm(1.25)
                s.paragraph_format.space_before = Pt(0)
                s.paragraph_format.space_after = Pt(0)
                s.paragraph_format.line_spacing = 1.5
                s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Табуляторы через безопасное API python-docx
                from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
                
                # Очищаем старые табы, если есть
                for ts in s.paragraph_format.tab_stops:
                    s.paragraph_format.tab_stops[0].clear()
                    
                s.paragraph_format.tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                
            except Exception as e:
                logger.error(f"Failed to style {style_name}: {e}")

    def add_title_page(
        self,
        university: str | None = None,
        work_type: str = "КУРСОВАЯ РАБОТА",
        topic: str = "",
        subject: str = "",
        student_name: str | None = None,
        student_group: str | None = None,
        teacher_name: str | None = None,
        teacher_title: str | None = None,
    ):
        """Добавить титульный лист."""
        def add_centered(text: str, size: int = 14, bold: bool = False, space_after: int = 0):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold = bold
            return p

        def add_right(text: str, size: int = 14):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            return p

        # Министерство
        add_centered("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ", size=12, space_after=0)
        add_centered("РОССИЙСКОЙ ФЕДЕРАЦИИ", size=12, space_after=0)

        # Университет
        if university:
            add_centered(university.upper(), size=12, bold=True, space_after=0)
        else:
            add_centered("НАИМЕНОВАНИЕ УНИВЕРСИТЕТА", size=12, bold=True, space_after=0)

        # Капитализация темы и очистка
        topic = topic.strip()
        if topic:
            topic = topic.strip('"').strip("'").strip('«').strip('»')
            topic = topic[0].upper() + topic[1:]
        
        # Нормализация ООО и ВКонтакте
        topic = re.sub(r'(?i)ооо\s+([a-zа-я]+)', r'ООО «\1»', topic)
        topic = re.sub(r'(?i)вконтакте', 'ВКонтакте', topic)

        # Пустые строки
        for _ in range(3):
            add_centered("")

        # Тип работы
        add_centered(work_type.upper(), size=16, bold=True, space_after=0)

        # Тема
        add_centered(f'на тему: "{topic}"', size=14, space_after=0)

        # Предмет
        add_centered(f"по дисциплине: {subject}", size=14, space_after=0)

        # Пустые строки
        for _ in range(3):
            add_centered("")

        # Студент и преподаватель (справа)
        if student_name:
            add_right(f"Выполнил(а): {student_name}")
        if student_group:
            add_right(f"Группа: {student_group}")
        if teacher_name:
            title = teacher_title or "Преподаватель"
            add_right(f"{title}: {teacher_name}")

        # Пустые строки до города
        for _ in range(4):
            add_centered("")

        # Город и год
        current_year = datetime.now().year
        add_centered(f"Москва, {current_year}", size=14)

        # Разрыв страницы (прикрепляем к последнему абзацу)
        if self.doc.paragraphs:
            self.doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
        else:
            self.doc.add_page_break()

    def add_table_of_contents(self):
        """Добавить автоматически обновляемое оглавление Word."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("СОДЕРЖАНИЕ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.bold = True

        # OXML для авто-оглавления
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run()
        fldChar = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        run._element.append(fldChar)
        
        run = paragraph.add_run()
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>' % nsdecls('w'))
        run._element.append(instrText)
        
        run = paragraph.add_run()
        fldChar = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
        run._element.append(fldChar)
        
        # Заглушка (Word обновит её при открытии)
        run = paragraph.add_run("Оглавление генерируется автоматически... (если не появилось — нажмите правой кнопкой и выберите «Обновить поле»)")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.italic = True
        
        run = paragraph.add_run()
        fldChar = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        run._element.append(fldChar)

    def add_heading_chapter(self, number: str, title: str):
        """Добавление заголовка главы (ГЛАВА 1. НАЗВАНИЕ) с разрывом страницы."""
        if self.doc.paragraphs:
            self.doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
        else:
            self.doc.add_page_break()
        h = self.doc.add_heading(level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_clean = self._sanitize_text(title.strip().upper())
        run = h.add_run(f"ГЛАВА {number}. {title_clean}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = True
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(0)

    def add_heading_section(self, number: str, title: str):
        """Добавление заголовка параграфа (1.1. НАЗВАНИЕ)."""
        h = self.doc.add_heading(level=2)
        # Подглавы (параграфы) НЕ КАПСОМ, а в формате "Первая буква заглавная"
        raw_title = title.strip()
        if raw_title:
            clean_title = raw_title[0].upper() + raw_title[1:].lower()
        else:
            clean_title = ""
        
        clean_title = self._sanitize_text(clean_title)
        
        run = h.add_run(f"{number} {clean_title}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = True
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(0)
        h.paragraph_format.first_line_indent = Cm(1.25)
        h.paragraph_format.keep_with_next = True  # ЗАГОЛОВОК НЕ ОТРЫВАЕТСЯ ОТ ТЕКСТА
        h.paragraph_format.keep_together = True

    def add_section_title(self, title: str, centered: bool = True, page_break: bool = True):
        """Добавить заголовок раздела (Введение, Заключение и т.д.)."""
        if page_break:
            if self.doc.paragraphs:
                self.doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
            else:
                self.doc.add_page_break()
        heading = self.doc.add_heading(level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        
        # Введение, Заключение и т.д. — ВСЕГДА КАПСОМ
        clean_title = self._sanitize_text(title.strip().upper())
        run = heading.add_run(clean_title)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.first_line_indent = Cm(0)
        heading.paragraph_format.keep_with_next = True

    def set_sources(self, sources: list[SourceItem]):
        """Установить список источников для генерации сносок."""
        self.sources_dict = {str(s.number): s for s in sources}

    def _fix_latex_escapes(self, text: str) -> str:
        """
        Восстанавливает LaTeX команды, которые могли быть интерпретированы Python 
        как управляющие символы (например, \f -> \x0c).
        """
        if not text:
            return ""
        # \f (form feed) -> \f
        text = text.replace('\x0c', '\\f')
        # \t (tab) -> \t (если это не намеренный таб, а часть команды типа \theta)
        # Но осторожно, таб может быть и в тексте. В формулах табов быть не должно.
        text = text.replace('\t', '\\t')
        # \v (vertical tab) -> \v
        text = text.replace('\x0b', '\\v')
        # \a (bell) -> \a
        text = text.replace('\x07', '\\a')
        # \b (backspace) -> \b
        text = text.replace('\x08', '\\b')
        
        return text

    def _sanitize_text(self, text: str) -> str:
        """Очистка текста: Markdown, визуальные маркеры, кавычки."""
        if not text:
            return ""

        # 0. КРИТИЧНО: Удаляем невалидные XML управляющие символы, которые ломают MS Word
        # (все символы с кодами от 0x00 до 0x1F, кроме \t, \n, \r)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)

        # 1. Удаляем плейсхолдеры визуала [ВСТАВИТЬ_ГРАФИК_N] / [ВСТАВИТЬ_ТАБЛИЦУ_N]
        text = re.sub(r'\[ВСТАВИТЬ_(ГРАФИК|ТАБЛИЦУ)_\d+\]', '', text)

        # 2. Удаляем Markdown-заголовки (# ..., ## ... и т.д.)
        text = re.sub(r'^#+\s+.*$', '', text, flags=re.MULTILINE)

        # 3. Удаляем Markdown жирный (**text**) и курсив (*text*), оставляя текст
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text, flags=re.DOTALL)
        text = re.sub(r'\*(.+?)\*', r'\1', text, flags=re.DOTALL)

        # 3.5 Защита многострочных формул и исправление экранирования
        def process_equation_match(match):
            eq_content = match.group(0)
            # Убираем лишние переносы
            eq_content = eq_content.replace('\n', ' ')
            # Чиним \frac (превращаем \x0c обратно в \f)
            eq_content = self._fix_latex_escapes(eq_content)
            return eq_content
            
        text = re.sub(r'\[EQUATION\].*?\[/EQUATION\]', process_equation_match, text, flags=re.DOTALL)

        # 4. Нормализуем более двух пустых строк до одного двойного переноса
        text = re.sub(r'\n{3,}', '\n\n', text)

        # 5. Заменяем английские кавычки на русские «ёлочки»
        text = text.replace('\u201c', '«').replace('\u201d', '»')
        text = re.sub(r'(\s|^)"', r'\1«', text)
        text = re.sub(r'"([\s.,!?;:)]|$)', r'»\1', text)

        # 6. Механическая замена тире: любые длинные тире и тире между пробелами -> среднее тире –
        text = text.replace('\u2014', '\u2013')  # em-dash -> en-dash
        text = text.replace(' — ', ' – ')        # на всякий случай явно
        text = text.replace(' - ', ' – ')
        text = text.replace(' ― ', ' – ')

        return text.strip()

    def add_text(self, text: str, style: str = "Normal"):
        """Добавление блока текста. Разбивка по \\n\\n = граница абзаца."""
        sanitized = self._sanitize_text(text)
        if not sanitized:
            return

        # Двойной перенос строки — граница абзаца (как задаёт LLM)
        raw_paragraphs = sanitized.strip().split('\n\n')
        for raw_para in raw_paragraphs:
            # Одиночные переносы внутри абзаца склеиваем в одну строку
            paragraph_text = ' '.join(
                line.strip() for line in raw_para.split('\n') if line.strip()
            )
            if not paragraph_text:
                continue
            p = self.doc.add_paragraph(style=style)
            
            # Проверка на примечание/источник или пояснение к формуле (где...)
            is_source_note = paragraph_text.startswith("Источник:")
            is_explanation = paragraph_text.lower().startswith("где")
            
            if is_source_note or is_explanation:
                p.paragraph_format.first_line_indent = Cm(0)  # Убираем отступ по ГОСТу
            else:
                p.paragraph_format.first_line_indent = Cm(1.25)
                
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            # Принудительно убираем интервалы для всех абзацев в подпунктах и введении
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            self._add_text_with_footnotes(p, paragraph_text)
            
            if is_source_note:
                for run in p.runs:
                    run.font.size = Pt(12)  # Уменьшаем шрифт до 12 pt

    def _add_text_with_footnotes(self, paragraph, text: str):
        """Парсинг маркеров [N], [N, M] и формул в тексте."""
        # Паттерн для сносок: ловим [1] и убираем съедание ведущего пробела перед скобкой
        footnote_pattern = r'\[(?P<ids>\d+(?:\s*,\s*\d+)*)(?:,\s*с\.\s*(?P<page>\d+(?:-\d+)?))?\]'
        eq_pattern = r'\[EQUATION\](?P<eq1>.*?)\[/EQUATION\]'
        hyperlink_pattern = r'\[HYPERLINK:(?P<url>.*?)\]'
        
        combined_pattern = f'(?P<footnote>{footnote_pattern})|(?P<equation>{eq_pattern})|(?P<hyperlink>{hyperlink_pattern})'
        
        last_idx = 0
        for match in re.finditer(combined_pattern, text, re.DOTALL):
            before_text = text[last_idx:match.start()]
            
            if match.group('footnote'):
                before_text = before_text.rstrip(' ')
            
            paragraph.add_run(before_text)

            if match.group('equation'):
                eq_text = match.group('eq1')
                self._insert_equation(paragraph, eq_text, (" [EQUATION]", "[/EQUATION] "))
            elif match.group('footnote'):
                source_ids = [s.strip() for s in match.group('ids').split(',')]
                page_info = match.group('page')
                
                for i, source_num in enumerate(source_ids):
                    source_obj = self.sources_dict.get(source_num)
                    if source_obj:
                        citation = source_obj.citation or source_obj.title or ""
                        url = source_obj.url
                    else:
                        citation = f"Источник №{source_num}"
                        url = None
                    
                    citation = citation.rstrip('. ')
                    # Если источников несколько, страницу пишем только в последнем или в каждом? 
                    # По ГОСТу лучше в каждом, если это разные книги.
                    footnote_text = f"{citation}. — С. {page_info}." if page_info else f"{citation}."
                    self._insert_footnote(paragraph, footnote_text, url)
                
            elif match.group('hyperlink'):
                url = match.group('url').strip()
                self._add_hyperlink(paragraph, url, url)
                
            last_idx = match.end()
        
        # Остаток текста
        paragraph.add_run(text[last_idx:])

        # Если параграф состоит ТОЛЬКО из уравнений (одного или нескольких), нумеруем и центрируем
        # Но по инструкции LLM должна быть одна формула на строку.
        stripped_text = text.strip()
        # Ищем все вхождения [EQUATION]...[/EQUATION] и убираем их
        text_without_equations = re.sub(r'\[EQUATION\].*?\[/EQUATION\]', '', stripped_text).strip()
        
        # Если в параграфе есть [EQUATION] и нет другого значимого текста
        if "[EQUATION]" in stripped_text and not text_without_equations:
            # Настраиваем табы для нумерации по ГОСТу
            # Центр (8.5 см) и Правый край (17.0 см)
            from docx.enum.text import WD_TAB_ALIGNMENT
            tab_stops = paragraph.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(8.5), WD_TAB_ALIGNMENT.CENTER)
            tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
            
            # Добавляем номер формулы в конце через табуляцию
            self.equation_counter += 1
            paragraph.add_run(f"\t({self.equation_counter})")
            paragraph.paragraph_format.first_line_indent = Cm(0) # У формул нет абзацного отступа
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT # Используем табы для позиционирования

    def _insert_equation(self, paragraph, latex_text: str, fallback_delims=None):
        """Конвертация LaTeX в OMML и вставка в параграф."""
        try:
            # --- ЗАЩИТА СИНТАКСИСА LATEX ---
            latex_safe = re.sub(r'(?<!\\)%', r'\%', latex_text)
            
            # 1. LaTeX -> MathML
            mathml = latex2mathml.converter.convert(latex_safe)
            # 2. MathML -> OMML
            omml = mathml2omml.convert(mathml)
            
            # 2.5 Исправление отсутствующего пространства имен m
            if 'xmlns:m' not in omml:
                omml = omml.replace('<m:oMath', '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"')
            
            # 3. Вставка XML в параграф
            omml_el = etree.fromstring(omml)
            
            # --- ИНЪЕКЦИЯ КИРИЛЛИЦЫ ---
            self._fix_omml_cyrillic(omml_el)
            
            # Если это отдельная строка для формулы, добавляем таб для центрирования
            if paragraph.text.startswith('\t') or not paragraph.text.strip():
                  paragraph.add_run('\t')
                  
            paragraph._element.append(omml_el)
        except Exception as e:
            logger.error(f"Error inserting equation '{latex_text}': {e}")
            # В случае ошибки вставляем как текст с исходными делемитерами
            d1, d2 = fallback_delims if fallback_delims else ("", "")
            paragraph.add_run(f" {d1}{latex_text}{d2} ")

    def _fix_omml_cyrillic(self, element):
        """Рекурсивно лечит кириллицу в OMML, задавая m:nor и Cambria Math."""
        M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        for t_node in element.xpath('.//m:t', namespaces={'m': M}):
            text = t_node.text or ""
            if re.search(r'[а-яА-ЯёЁ]', text):
                r_node = t_node.getparent()
                if r_node is not None:
                    rPr = r_node.find('m:rPr', namespaces={'m': M})
                    if rPr is None:
                        rPr = etree.Element('{%s}rPr' % M)
                        r_node.insert(0, rPr)
                    if rPr.find('m:nor', namespaces={'m': M}) is None:
                        etree.SubElement(rPr, '{%s}nor' % M)
                    rFonts = rPr.find('m:rFonts', namespaces={'m': M})
                    if rFonts is None:
                        rFonts = etree.SubElement(rPr, '{%s}rFonts' % M)
                    rFonts.set('{%s}ascii' % M, 'Cambria Math')
                    rFonts.set('{%s}hAnsi' % M, 'Cambria Math')
                    rFonts.set('{%s}cs' % M, 'Cambria Math')
        for child in element:
            self._fix_omml_cyrillic(child)

    def _insert_footnote(self, paragraph, footnote_text: str, url: str | None = None):
        """Вставка настоящей нижней сноски Word с поддержкой активных ссылок и 10pt шрифтом."""
        # Простая валидация URL: должен содержать точку и начинаться с http/https
        is_valid_url = False
        if url and isinstance(url, str):
            url = url.strip()
            if url.startswith(('http://', 'https://')) and '.' in url.split('//')[-1]:
                is_valid_url = True
        
        # Если URL не валиден, зануляем его
        if not is_valid_url:
            url = None

        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

        if self._footnotes_part is None:
            self._footnote_id += 1
            run = paragraph.add_run(f"[{self._footnote_id}]")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.superscript = True
            return

        self._footnote_id += 1
        fid = self._footnote_id
        
        # Санитизируем текст сноски (включая тире)
        footnote_text = self._sanitize_text(footnote_text)

        # --- 1. Добавляем запись в footnotes.xml ---
        try:
            fn_root = etree.fromstring(self._footnotes_part.blob)
            fn_el = etree.SubElement(fn_root, '{%s}footnote' % W)
            fn_el.set('{%s}id' % W, str(fid))

            fn_p = etree.SubElement(fn_el, '{%s}p' % W)
            fn_pPr = etree.SubElement(fn_p, '{%s}pPr' % W)
            fn_pStyle = etree.SubElement(fn_pPr, '{%s}pStyle' % W)
            fn_pStyle.set('{%s}val' % W, 'CalamoFootnoteText')
            
            # Абзацный отступ (first line indent) 1.25 см (709 twips)
            fn_ind = etree.SubElement(fn_pPr, '{%s}ind' % W)
            fn_ind.set('{%s}left' % W, '0')
            fn_ind.set('{%s}firstLine' % W, '709')

            # Выравнивание по ширине
            fn_jc = etree.SubElement(fn_pPr, '{%s}jc' % W)
            fn_jc.set('{%s}val' % W, 'both')

            # Номер сноски (автоматический)
            fn_r_ref = etree.SubElement(fn_p, '{%s}r' % W)
            fn_rPr_ref = etree.SubElement(fn_r_ref, '{%s}rPr' % W)
            fn_rStyle = etree.SubElement(fn_rPr_ref, '{%s}rStyle' % W)
            fn_rStyle.set('{%s}val' % W, 'CalamoFootnoteReference')
            etree.SubElement(fn_r_ref, '{%s}footnoteRef' % W)

            # Текст сноски. Нормализуем строку ГОСТа перед вставкой.
            footnote_text = self._normalize_gost_citation(footnote_text)
            
            # Ищем тег [HYPERLINK:...]
            link_match = re.search(r'\[HYPERLINK:(?P<url>.*?)\]', footnote_text)
            
            if link_match:
                url_to_insert = link_match.group('url').strip()
                # Удаляем сам тег из текста, чтобы он не печатался
                display_text = footnote_text.replace(link_match.group(0), url_to_insert)
                parts = display_text.split(url_to_insert)
                
                # Часть ДО ссылки
                if parts[0]:
                    fn_r_pre = etree.SubElement(fn_p, '{%s}r' % W)
                    fn_rPr_pre = etree.SubElement(fn_r_pre, '{%s}rPr' % W)
                    etree.SubElement(fn_rPr_pre, '{%s}sz' % W).set('{%s}val' % W, '20')
                    fn_t_pre = etree.SubElement(fn_r_pre, '{%s}t' % W)
                    fn_t_pre.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    fn_t_pre.text = f' {parts[0]}'
                
                # Кликабельная ссылка
                r_id = self._footnotes_part.relate_to(url_to_insert, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
                fn_hyper = etree.SubElement(fn_p, '{%s}hyperlink' % W)
                fn_hyper.set('{%s}id' % R, r_id)
                fn_r_link = etree.SubElement(fn_hyper, '{%s}r' % W)
                fn_rPr_link = etree.SubElement(fn_r_link, '{%s}rPr' % W)
                fn_rStyle_link = etree.SubElement(fn_rPr_link, '{%s}rStyle' % W)
                fn_rStyle_link.set('{%s}val' % W, 'Hyperlink')
                etree.SubElement(fn_rPr_link, '{%s}sz' % W).set('{%s}val' % W, '20')
                fn_t_link = etree.SubElement(fn_r_link, '{%s}t' % W)
                fn_t_link.text = url_to_insert
                
                # Часть ПОСЛЕ ссылки
                if len(parts) > 1 and parts[1]:
                    fn_r_post = etree.SubElement(fn_p, '{%s}r' % W)
                    fn_rPr_post = etree.SubElement(fn_r_post, '{%s}rPr' % W)
                    etree.SubElement(fn_rPr_post, '{%s}sz' % W).set('{%s}val' % W, '20')
                    fn_t_post = etree.SubElement(fn_r_post, '{%s}t' % W)
                    fn_t_post.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    fn_t_post.text = parts[1]
            else:
                fn_r_txt = etree.SubElement(fn_p, '{%s}r' % W)
                fn_rPr_txt = etree.SubElement(fn_r_txt, '{%s}rPr' % W)
                etree.SubElement(fn_rPr_txt, '{%s}sz' % W).set('{%s}val' % W, '20')
                fn_t = etree.SubElement(fn_r_txt, '{%s}t' % W)
                fn_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                fn_t.text = f' {footnote_text}'

            self._footnotes_part._blob = etree.tostring(fn_root, xml_declaration=True, encoding='UTF-8', standalone=True)
        except Exception as e:
            logger.error(f"Error in _insert_footnote: {e}")
            pass

        # --- 2. Вставляем ссылку на сноску в параграф ---
        run = paragraph.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        rPr_xml = ('<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                   '<w:rStyle w:val="CalamoFootnoteReference"/>'
                   '<w:vertAlign w:val="superscript"/></w:rPr>')
        rPr_el = etree.fromstring(rPr_xml)
        run._element.insert(0, rPr_el)

        ref_xml = ('<w:footnoteReference xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                   f' w:id="{fid}"/>')
        run._element.append(etree.fromstring(ref_xml))

    def add_figure(self, image_path: Path, figure_number: int, title: str, source: str = "составлено автором", skip_header: bool = False):
        """Вставить рисунок с подписью по ГОСТу."""
        # Рисунок по центру
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        if image_path.exists():
            run = p.add_run()
            run.add_picture(str(image_path), width=Mm(170))

        if not skip_header:
            # Подпись под рисунком (Рисунок 1 – Название)
            caption = self.doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.first_line_indent = Cm(0)
            caption.paragraph_format.space_after = Pt(0)
            title_clean = self._sanitize_text(title)
            run = caption.add_run(f"Рисунок {figure_number} – {title_clean}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.bold = False
            # Источник под рисунком
            self._add_source_paragraph(source)

    def add_table_data(self, table_number: int, title: str, headers: list, rows: list, source: str = "составлено автором", skip_header: bool = False):
        """Добавление таблицы с подписью СВЕРХУ по ГОСТу."""
        if not skip_header:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            # Подпись таблицы (Таблица 1 – Название)
            title_clean = self._sanitize_text(title)
            run = p.add_run(f"Таблица {table_number} – {title_clean}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

        # Создаем таблицу
        if not rows: return
        
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заголовки
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = self._sanitize_text(str(h))
            para = hdr_cells[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            if para.runs:
                run = para.runs[0]
                run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        # Данные
        for i, row_data in enumerate(rows):
            row_cells = table.rows[i+1].cells
            for j, val in enumerate(row_data):
                clean_val = str(val).strip('|').strip()
                if clean_val == '-': clean_val = ""
                row_cells[j].text = self._sanitize_text(clean_val)
                
                para = row_cells[j].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                if para.runs:
                    run = para.runs[0]
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
        
        if not skip_header:
            self._add_source_paragraph(source)

    def _add_hyperlink(self, paragraph, url: str, text: str, font_size: int = 14):
        """Добавить кликабельную ссылку в абзац."""
        r_id = paragraph.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        sz_val = font_size * 2
        hyperlink = parse_xml(
            f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            f' r:id="{r_id}"><w:r><w:rPr><w:rStyle w:val="CalamoHyperlink"/>'
            f'<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            f'<w:sz w:val="{sz_val}"/></w:rPr><w:t>{text}</w:t></w:r></w:hyperlink>'
        )
        paragraph._element.append(hyperlink)

    def _normalize_gost_citation(self, text: str) -> str:
        """Алгоритмическая нормализация строки ГОСТ: перемещение страниц перед URL."""
        if not text: return ""
        
        # Паттерны для поиска URL и страниц
        url_pattern = r'–?\s*URL:\s*(https?://[^\s\)\],]+)'
        pages_pattern = r'–?\s*С\.\s*\d+[-–]\d+\.?'
        total_pages_pattern = r'–?\s*\d+\s*с\.?'
        
        url_match = re.search(url_pattern, text)
        if not url_match:
            return text
            
        url_start = url_match.start()
        after_url = text[url_start:]
        before_url = text[:url_start]
        
        # Ищем страницы ПОСЛЕ URL
        found_pages = re.search(pages_pattern, after_url) or re.search(total_pages_pattern, after_url)
        
        if found_pages:
            pages_str = found_pages.group(0).strip()
            # Удаляем страницы из хвоста
            after_url_clean = after_url.replace(found_pages.group(0), "").strip()
            # Формируем новую строку: [ДО] [СТРАНИЦЫ] [URL + ОСТАТОК]
            # Убеждаемся в наличии разделителей
            sep = " – "
            normalized = f"{before_url.rstrip('. –')}{sep}{pages_str.lstrip('– ')}{sep}{after_url_clean.lstrip('– ')}"
            # Чистим двойные тире
            normalized = re.sub(r'\s*[–—]\s*[–—]\s*', ' – ', normalized)
            return normalized.strip()
            
        return text

    def _add_source_paragraph(self, source_text: str):
        """Добавляет источник с кликабельными ссылками на новой строке (12pt)."""
        src_p = self.doc.add_paragraph()
        src_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        src_p.paragraph_format.first_line_indent = Cm(1.25)
        src_p.paragraph_format.space_before = Pt(0)
        src_p.paragraph_format.space_after = Pt(0)
        src_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        prefix = "Источник: "
        if source_text.lower().startswith("источник:"):
            source_text = source_text[9:].strip()
            
        source_text = self._sanitize_text(source_text)
        full_text = prefix + source_text
        # Ищем теги [HYPERLINK:...]
        import re
        link_pattern = re.compile(r'\[HYPERLINK:(?P<url>.*?)\]')
        
        last_idx = 0
        for match in link_pattern.finditer(full_text):
            # Текст ДО ссылки
            before = full_text[last_idx:match.start()]
            if before:
                run = src_p.add_run(before)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            
            # Сама ссылка
            url = match.group('url').strip()
            self._add_hyperlink(src_p, url, url, font_size=12)
            
            last_idx = match.end()
            
        # Остаток текста
        remaining = full_text[last_idx:]
        if remaining:
            run = src_p.add_run(remaining)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    def add_sources_list_gost(self, sources: list[SourceItem]):
        """Добавить список использованных источников с кликабельными ссылками и ГОСТ отступами."""
        for source in sources:
            p = self.doc.add_paragraph(style="Normal")
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            citation_text = source.citation or source.title or ""
            citation_text = self._sanitize_text(citation_text)
            # Ищем тег [HYPERLINK:...]
            import re
            link_match = re.search(r'\[HYPERLINK:(?P<url>.*?)\]', citation_text)
            
            if link_match:
                url_to_insert = link_match.group('url').strip()
                # Удаляем сам тег из текста, чтобы он не печатался
                display_text = citation_text.replace(link_match.group(0), url_to_insert)
                parts = display_text.split(url_to_insert)
                
                # Часть ДО ссылки (включая номер)
                run = p.add_run(f"{source.number}. {parts[0]}")
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                
                # Кликабельная ссылка
                try:
                    self._add_hyperlink(p, url_to_insert, url_to_insert, font_size=14)
                except Exception:
                    run_url = p.add_run(url_to_insert)
                    run_url.font.name = "Times New Roman"
                    run_url.font.size = Pt(14)
                
                # Часть ПОСЛЕ ссылки
                if len(parts) > 1 and parts[1]:
                    run_post = p.add_run(parts[1])
                    run_post.font.name = "Times New Roman"
                    run_post.font.size = Pt(14)
            else:
                full_text = f"{source.number}. {citation_text}"
                run = p.add_run(full_text)
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)

    def add_page_numbers(self):
        """Нумерация страниц уже добавлена в _setup_page_layout. Метод оставлен для совместимости."""
        pass

    def save(self, filepath: Path) -> Path:
        """Сохранить документ."""
        filepath.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(filepath))
        return filepath


def create_docx_builder() -> DocxBuilder:
    """Фабрика для создания билдера."""
    return DocxBuilder()
