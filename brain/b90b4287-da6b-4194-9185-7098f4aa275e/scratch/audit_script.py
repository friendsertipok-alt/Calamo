# -*- coding: utf-8 -*-
import docx
import os
import json

def analyze_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    tables_count = len(doc.tables)
    images_count = 0
    
    # Считаем изображения (приблизительно через inline_shapes)
    images_count = len(doc.inline_shapes)
    
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # Анализ на наличие Markdown
    markdown_elements = []
    for p in paragraphs:
        if "#" in p or "**" in p or "__" in p or "###" in p:
            markdown_elements.append(p)
            
    # Анализ оформления названий (Sentence case)
    bad_titles = []
    for p in paragraphs:
        if p.isupper() and len(p) > 10: # Подозрение на капс в названии
            bad_titles.append(p)
            
    # Анализ объема анализа после визуалов
    # (Это сложнее, просто соберем текст для LLM анализа)
    
    analysis_report = {
        "filename": os.path.basename(file_path),
        "total_paragraphs": len(paragraphs),
        "tables_found": tables_count,
        "images_found": images_count,
        "markdown_issues": markdown_elements[:5], # Примеры
        "caps_titles": bad_titles[:5],
        "full_content": "\n".join(paragraphs[:100]) # Первые 100 абзацев для контекста
    }
    
    print(json.dumps(analysis_report, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    analyze_docx(r"c:\Users\sevam\OneDrive\Рабочий стол\folders\antigravity\Calamo\Academicpro\analysis_room\Paper_f2960ab9.docx")
